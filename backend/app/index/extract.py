"""Document text extraction — REQ-16.

Returns text split into labelled sections rather than one flat blob, because the
label is what makes a citation useful. "It's in lease.pdf" sends someone hunting
through 40 pages; "lease.pdf, page 12" does not.

Extraction runs locally. Nothing here reaches the network.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

log = logging.getLogger(__name__)

SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".markdown", ".log", ".csv"}


class ExtractionError(Exception):
    """The file could not be read. Recorded against the document, not raised on."""


@dataclass
class Section:
    label: str  # "page 12", "Termination clause", or "" for a whole small file
    text: str


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED


def extract(path: Path) -> list[Section]:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(path)
        if suffix == ".docx":
            return _extract_docx(path)
        return _extract_text(path)
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 — one unreadable file must not stop a scan
        raise ExtractionError(str(exc)) from exc


def _extract_pdf(path: Path) -> list[Section]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionError("pypdf is not installed") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"unreadable PDF: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        # Try the empty password, which covers PDFs encrypted only against
        # editing. A genuinely locked file is reported, never guessed at.
        try:
            if reader.decrypt("") == 0:
                raise ExtractionError("password protected")
        except ExtractionError:
            raise
        except Exception as exc:  # noqa: BLE001
            raise ExtractionError("password protected") from exc

    sections: list[Section] = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 — a bad page shouldn't lose the document
            log.debug("page %d of %s failed to extract", number, path)
            continue
        text = _tidy(text)
        if text:
            sections.append(Section(f"page {number}", text))

    if not sections:
        # Almost always a scanned document. Say so plainly — silently indexing
        # nothing would leave the user thinking the file was searched.
        raise ExtractionError("no extractable text (likely a scan; OCR not supported yet)")
    return sections


def _extract_docx(path: Path) -> list[Section]:
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError("python-docx is not installed") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"unreadable DOCX: {exc}") from exc

    sections: list[Section] = []
    heading = ""
    buffer: list[str] = []

    def flush() -> None:
        body = _tidy("\n".join(buffer))
        if body:
            sections.append(Section(heading, body))
        buffer.clear()

    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        # Word's built-in heading styles are the document's own structure —
        # far better section labels than an arbitrary character offset.
        style = paragraph.style
        if (getattr(style, "name", None) or "").startswith("Heading"):
            flush()
            heading = text
            continue
        buffer.append(text)
    flush()

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            sections.append(Section(heading or "table", "\n".join(rows)))

    if not sections:
        raise ExtractionError("document contains no text")
    return sections


_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")


def _extract_text(path: Path) -> list[Section]:
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            raw = path.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            raise ExtractionError(str(exc)) from exc
    except OSError as exc:
        raise ExtractionError(str(exc)) from exc

    if path.suffix.lower() not in {".md", ".markdown"}:
        text = _tidy(raw)
        if not text:
            raise ExtractionError("file is empty")
        return [Section("", text)]

    sections: list[Section] = []
    heading = ""
    buffer: list[str] = []

    def flush() -> None:
        body = _tidy("\n".join(buffer))
        if body:
            sections.append(Section(heading, body))
        buffer.clear()

    for line in raw.splitlines():
        match = _MD_HEADING.match(line)
        if match:
            flush()
            heading = match.group(2)
            continue
        buffer.append(line)
    flush()

    if not sections:
        raise ExtractionError("file is empty")
    return sections


def title_for(path: Path, sections: list[Section]) -> str:
    """A human-facing name for the document, for use in citations."""
    for section in sections:
        if section.label and not section.label.startswith("page "):
            return section.label
    return path.stem


def _tidy(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
